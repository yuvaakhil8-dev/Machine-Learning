from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "outputs" / "reports" / "project_work_done"
OUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT_DIR / "Safe_Unsafe_Video_Classification_Work_Done_Report.docx"
MD_PATH = OUT_DIR / "Safe_Unsafe_Video_Classification_Work_Done_Report.md"


def read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", newline="", encoding="utf-8-sig") as file:
        return list(csv.DictReader(file))


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, value: object, *, bold: bool = False, size: float = 8.5, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(value))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[object]], *, size: float = 8.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell_text(table.rows[0].cells[idx], header, bold=True, size=size, color="FFFFFF")
        shade(table.rows[0].cells[idx], "1F4E79")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell_text(cells[idx], value, size=size)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


quality = json.loads((ROOT / "outputs" / "audit" / "dataset_audit" / "video_quality_summary.json").read_text(encoding="utf-8"))
metrics = read_csv(ROOT / "outputs" / "reports" / "project_submission_evidence" / "model_metrics_clean.csv")
cv = read_csv(ROOT / "models_artifacts" / "artifacts_phase1" / "cross_validation_results.csv")
feature_rows = read_csv(ROOT / "data" / "features" / "video_features_advanced.csv")
best_metric = max(metrics, key=lambda row: float(row["test_accuracy"]))
best_cv = max(cv, key=lambda row: float(row["cv_accuracy_mean"]))

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.6)
section.right_margin = Inches(0.6)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SAFE AND UNSAFE VIDEO CLASSIFICATION")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(31, 78, 121)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Work Done Report")
r.bold = True
r.font.size = Pt(14)

doc.add_paragraph(
    "This report summarizes the completed work in the project, including dataset preparation, feature extraction, "
    "model training, stacking classifier implementation, cross-validation, result analysis, web app development, "
    "project organization, and final submission outputs."
)

doc.add_heading("1. Project Objective", level=1)
doc.add_paragraph(
    "The project classifies videos into Safe and Unsafe categories using computer vision and machine learning. "
    "The final system extracts frame-based visual features, motion/activity features, object/person cues, proxy activity cues, "
    "and CNN embedding features, then compares multiple machine-learning models and provides a Streamlit web interface for demo prediction."
)

doc.add_heading("2. Dataset Work Completed", level=1)
add_table(
    doc,
    ["Item", "Final Output"],
    [
        ["Total valid videos", quality["total_videos"]],
        ["Safe videos", quality["class_counts"]["safe"]],
        ["Unsafe videos", quality["class_counts"]["unsafe"]],
        ["Worthy videos", quality["status_counts"]["worthy"]],
        ["Videos marked for review", quality["status_counts"]["review"]],
        ["Cannot-open videos", quality["cannot_open_count"]],
        ["Feature CSV rows", len(feature_rows)],
        ["Feature CSV columns", len(feature_rows[0]) if feature_rows else 0],
    ],
)
add_bullets(
    doc,
    [
        "Original video folders were preserved.",
        "Dataset audit was generated to check duration, resolution, openability, and class distribution.",
        "Empty non-video files were excluded from the clean submission package.",
        "A clean submission zip was created with dataset/safe and dataset/unsafe folders.",
    ],
)

doc.add_heading("3. Feature Extraction Work Completed", level=1)
add_table(
    doc,
    ["Feature group", "What was implemented", "Output"],
    [
        ["CNN embeddings", "Sampled frames are passed through a pretrained CNN and averaged into video-level features.", "512 CNN feature columns"],
        ["Motion/activity", "Frame difference, optical-flow magnitude, acceleration, and temporal consistency are calculated.", "Motion and activity scores"],
        ["YOLO/object cues", "Person/object/crowd/suspicious-object features are extracted as interpretable visual cues.", "YOLO feature columns"],
        ["Face/activity proxies", "Lightweight face count, expression proxy, pose velocity, and movement instability features are calculated.", "Proxy activity columns"],
        ["Fusion features", "Multiple visual signals are combined into final video activity and risk proxy scores.", "Fusion score columns"],
    ],
)
doc.add_paragraph(
    "The project uses more than 100 features; the final feature table contains 519 columns including metadata, labels, and visual descriptors."
)

doc.add_heading("4. Models Implemented", level=1)
add_bullets(
    doc,
    [
        "Logistic Regression",
        "Random Forest",
        "Support Vector Machine",
        "XGBoost",
        "KNN experiment",
        "Stacking Classifier",
        "CNN-LSTM module/design for temporal deep-learning extension",
    ],
)

doc.add_heading("5. Stacking Classifier Work", level=1)
doc.add_paragraph(
    "The stacking classifier combines multiple base learners and uses a meta-learner to produce the final Safe/Unsafe prediction. "
    "Instead of depending on one model, it combines model outputs to improve robustness. The saved stacking model, predictions, and confusion matrix are stored in the model artifacts folder."
)

doc.add_heading("6. Result Analysis", level=1)
add_table(
    doc,
    ["Model", "Test Accuracy", "Precision", "Recall", "F1", "ROC-AUC", "Status"],
    [
        [
            row["phase"],
            row["test_accuracy"],
            row["precision"],
            row["recall"],
            row["f1_score"],
            row["roc_auc"],
            row.get("generalization_status", "CV validated"),
        ]
        for row in metrics
    ],
)
doc.add_paragraph(
    f"Best test model: {best_metric['phase']} with {float(best_metric['test_accuracy']):.2%} test accuracy, "
    f"{float(best_metric['f1_score']):.2%} F1 score, and {float(best_metric['roc_auc']):.2%} ROC-AUC."
)

doc.add_heading("7. Cross-Validation Work", level=1)
add_table(
    doc,
    ["Model", "CV Accuracy", "CV F1", "CV ROC-AUC"],
    [[row["phase"], row["cv_accuracy_mean"], row["cv_f1_mean"], row["cv_roc_auc_mean"]] for row in cv],
)
doc.add_paragraph(
    f"Best cross-validation model: {best_cv['phase']} with {float(best_cv['cv_accuracy_mean']):.2%} CV accuracy, "
    f"{float(best_cv['cv_f1_mean']):.2%} CV F1 score, and {float(best_cv['cv_roc_auc_mean']):.2%} CV ROC-AUC. "
    "The final status is reported as CV validated."
)

doc.add_heading("8. Web App Work Completed", level=1)
add_bullets(
    doc,
    [
        "Built a Streamlit web app with top tabs: Overview, Dataset, Pipeline, Results, and Demo.",
        "Added detailed Pipeline section where each block explains input, internal process, output, and output files.",
        "Added live Demo tab for uploading a video and showing Safe/Unsafe prediction.",
        "Added readable feature display instead of raw visual_000 style labels.",
        "Added caching so repeated outputs load faster.",
    ],
)
doc.add_paragraph("Current web link: http://localhost:8501")

doc.add_heading("9. Project Organization Work", level=1)
add_table(
    doc,
    ["Folder", "Purpose"],
    [
        ["app/", "Streamlit web application"],
        ["src/", "Source code, feature engineering, models, preprocessing, evaluation, and scripts"],
        ["data/", "Feature CSVs, metadata, feature references, and model weights"],
        ["models_artifacts/", "Trained models, metrics, predictions, confusion matrices, and cross-validation outputs"],
        ["outputs/", "Reports, PPTs, audit files, and evidence documents"],
        ["submission/", "Final clean submission zip"],
        ["Safe (3)/ and Unsafe/", "Original video dataset folders"],
    ],
)

doc.add_heading("10. Final Submission Outputs", level=1)
add_bullets(
    doc,
    [
        "Clean submission zip: submission/Safe_Unsafe_Video_Classification_Submission.zip",
        "Full project report: outputs/reports/final_report/Safe_Unsafe_Video_Classification_Full_Project_Report.docx",
        "Work done report: outputs/reports/project_work_done/Safe_Unsafe_Video_Classification_Work_Done_Report.docx",
        "IEEE feature reference audit: outputs/reports/project_submission_evidence/IEEE_Feature_References_and_Project_Audit.docx",
        "Presentation files: outputs/presentations/final_ppt/",
        "Clean metrics file: outputs/reports/project_submission_evidence/model_metrics_clean.csv",
    ],
)

doc.add_heading("11. Final Conclusion", level=1)
doc.add_paragraph(
    "The project now has a complete video-only machine-learning pipeline: dataset audit, feature extraction, descriptive feature files, "
    "multiple model training, stacking classifier, cross-validation, result visualization, web app demo, IEEE-style feature justification, "
    "clean folder organization, and final submission package. The project is ready to explain as a complete end-to-end Safe/Unsafe video classification system."
)

doc.save(DOCX_PATH)

md = f"""# Safe and Unsafe Video Classification - Work Done Report

## Main Work Completed
- Dataset audited: {quality['total_videos']} valid videos.
- Safe videos: {quality['class_counts']['safe']}; Unsafe videos: {quality['class_counts']['unsafe']}.
- Feature table created: {len(feature_rows)} rows and {len(feature_rows[0]) if feature_rows else 0} columns.
- Models trained: Logistic Regression, Random Forest, SVM, XGBoost, KNN, Stacking Classifier.
- Best model: {best_metric['phase']} with {float(best_metric['test_accuracy']):.2%} test accuracy.
- Best cross-validation: {best_cv['phase']} with {float(best_cv['cv_accuracy_mean']):.2%} CV accuracy.
- Web app created at http://localhost:8501.
- Final submission zip created in `submission/`.

## Organized Folder Structure
- `app/` - Streamlit app
- `src/` - source code and scripts
- `data/` - feature files, metadata, references, weights
- `models_artifacts/` - trained models and metrics
- `outputs/` - reports, PPT, audit, evidence
- `submission/` - final zip
- `Safe (3)/`, `Unsafe/` - original dataset videos
"""
MD_PATH.write_text(md, encoding="utf-8")

print(DOCX_PATH)
print(MD_PATH)
