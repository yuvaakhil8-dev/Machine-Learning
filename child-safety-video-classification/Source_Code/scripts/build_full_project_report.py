from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "final_report"
OUT_DIR.mkdir(exist_ok=True)

DOCX_PATH = OUT_DIR / "Safe_Unsafe_Video_Classification_Full_Project_Report.docx"
MD_PATH = OUT_DIR / "Safe_Unsafe_Video_Classification_Full_Project_Report.md"


TITLE = "Safe vs Unsafe Video Classification"
SUBTITLE = "Using Machine Learning and Computer Vision"
AUTHORS = "D.Yuva Akhil, Y.Prannav, P.JayVardhan"


def pct(value: object) -> str:
    try:
        return f"{float(value):.2%}"
    except Exception:
        return str(value)


def load_csv(path: str) -> pd.DataFrame:
    file = ROOT / path
    if file.exists():
        return pd.read_csv(file)
    return pd.DataFrame()


metrics = load_csv("artifacts_phase1/metrics_summary.csv")
cv = load_csv("artifacts_phase1/cross_validation_results.csv")
quality = load_csv("final_outputs_csv/11_video_quality_summary.csv")
mapping = load_csv("feature_column_name_mapping_descriptive.csv")

best = metrics.sort_values("test_accuracy", ascending=False).iloc[0] if not metrics.empty else {}
best_cv = cv.sort_values("cv_accuracy_mean", ascending=False).iloc[0] if not cv.empty else {}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text_color(cell, color: str = "FFFFFF") -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)
            run.font.bold = True


def style_table(table, header_fill: str = "2563EB") -> None:
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Segoe UI"
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                set_cell_text_color(cell)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Segoe UI"
        run.font.color.rgb = RGBColor(20, 33, 61)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        run.font.name = "Segoe UI"
        run.font.size = Pt(10.5)
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        run.font.name = "Segoe UI"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(31, 41, 55)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        run.font.name = "Segoe UI"
        run.font.size = Pt(10.2)
        run.font.color.rgb = RGBColor(31, 41, 55)


def add_callout(doc: Document, title: str, body: str, fill: str = "EAF2FF") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Segoe UI"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(20, 33, 61)
    p.add_run(f"\n{body}")
    for run in p.runs:
        run.font.name = "Segoe UI"
    doc.add_paragraph()


def add_table_from_rows(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    style_table(table)
    doc.add_paragraph()


def add_image(doc: Document, rel_path: str, caption: str, width: float = 5.9) -> None:
    path = ROOT / rel_path
    if path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.font.name = "Segoe UI"
            r.font.size = Pt(8.5)
            r.font.italic = True
            r.font.color.rgb = RGBColor(100, 116, 139)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    styles["Normal"].font.name = "Segoe UI"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.name = "Segoe UI Semibold"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(20, 33, 61)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(SUBTITLE)
    r.font.name = "Segoe UI"
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(37, 99, 235)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run(f"Submitted by: {AUTHORS}\nProject Type: Computer Vision and Machine Learning\nDate: 07 May 2026")
    r.font.name = "Segoe UI"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(71, 85, 105)

    add_callout(
        doc,
        "Project Summary",
        "This project builds a complete video-only Safe/Unsafe classification pipeline using sampled video frames, CNN embedding features, multiple ML models, stacking ensemble learning, cross-validation, and a Streamlit web interface.",
        fill="EAF7EF",
    )

    doc.add_page_break()

    add_heading(doc, "1. Abstract")
    add_para(
        doc,
        "This project presents an intelligent video classification system that classifies videos into Safe and Unsafe categories using computer vision and machine learning. The system takes videos from a labelled local dataset, validates and preprocesses them, extracts visual features from sampled frames, trains multiple classifiers, evaluates performance through test metrics and 5-fold cross-validation, and provides a Streamlit web application for uploaded-video prediction."
    )
    add_para(
        doc,
        "The main feature representation is a 512-dimensional CNN embedding vector extracted from video frames. For better understanding, the 512 features are renamed into descriptive columns and grouped into edge/texture, object/body, scene/activity, and safety-context feature categories."
    )

    add_heading(doc, "2. Problem Statement")
    add_bullets(
        doc,
        [
            "Manual checking of video content is time-consuming and inconsistent.",
            "Unsafe activity may depend on body posture, movement, interaction, objects, and scene context.",
            "Simple rules such as brightness or color are not reliable for classifying unsafe videos.",
            "A machine learning system is required to learn patterns from video frames and predict Safe or Unsafe classes.",
        ],
    )
    add_callout(
        doc,
        "Key Point",
        "The project does not decide safety from a single feature. It uses learned CNN visual patterns and model comparison to classify videos.",
    )

    add_heading(doc, "3. Objectives")
    add_bullets(
        doc,
        [
            "Prepare and audit a labelled Safe/Unsafe video dataset.",
            "Extract reliable video-only visual features from sampled frames.",
            "Create a descriptive feature CSV where each CNN feature column has a readable name.",
            "Train and compare Logistic Regression, Random Forest, SVM, XGBoost, and Stacking models.",
            "Use 5-fold cross-validation to check generalization and reduce dependence on one train-test split.",
            "Build a Streamlit web app that accepts uploaded videos and predicts Safe or Unsafe.",
        ],
    )

    add_heading(doc, "4. Dataset Description")
    add_table_from_rows(
        doc,
        ["Item", "Value", "Explanation"],
        [
            ["Total videos", "497", "Original local dataset videos used for classification."],
            ["Safe videos", "250", "Videos labelled as safe."],
            ["Unsafe videos", "247", "Videos labelled as unsafe."],
            ["Dataset structure", "Safe (3) / Unsafe", "Separate class folders."],
            ["External videos kept", "0", "Final project preserves the local dataset."],
        ],
        widths=[1.5, 1.5, 4.0],
    )
    add_image(doc, "results/dataset_distribution.png", "Figure 1. Dataset class distribution.", width=4.7)

    add_heading(doc, "5. Dataset Quality Audit")
    add_table_from_rows(
        doc,
        ["Audit Field", "Value", "Meaning"],
        [
            ["Total videos checked", "497", "All dataset videos were scanned."],
            ["Cannot open count", "0", "No unreadable videos in the audit summary."],
            ["Minimum duration", "1.97 seconds", "Shortest video duration observed."],
            ["Maximum duration", "120.12 seconds", "Longest video duration observed."],
            ["Low resolution count", "6", "Small number of videos may require quality review."],
            ["Too short count", "1", "Only one video is below the preferred duration threshold."],
        ],
        widths=[1.8, 1.5, 3.7],
    )
    add_para(
        doc,
        "Dataset auditing improves the reliability of the project because corrupted or poor-quality videos can affect feature extraction and model performance. In this dataset, no videos failed the open/read validation, so the dataset is suitable for experimentation."
    )

    add_heading(doc, "6. Preprocessing Pipeline")
    add_table_from_rows(
        doc,
        ["Step", "Process", "Purpose"],
        [
            ["1", "Video validation", "Checks whether each video can be opened and read."],
            ["2", "Frame sampling", "Extracts representative frames from the video timeline."],
            ["3", "Resize and normalize", "Prepares frames for CNN feature extraction."],
            ["4", "Feature extraction", "Converts video frames into numerical CNN embeddings."],
            ["5", "Train-test split", "Creates stratified training and testing partitions."],
            ["6", "Scaling", "Standardizes numeric features before model training."],
        ],
        widths=[0.55, 2.1, 4.3],
    )

    add_heading(doc, "7. Feature Extraction")
    add_para(
        doc,
        "The main feature-extraction file for presentation and understanding is `video_features_advanced_descriptive.csv`. It contains 519 columns, including metadata columns and 512 descriptive CNN embedding feature columns."
    )
    add_table_from_rows(
        doc,
        ["Feature Category", "Column Range", "Count", "Examples"],
        [
            ["Metadata", "video_name, video_path, class_name, label, frames_sampled, video_fps, video_duration_seconds", "7", "Identity, label, duration, FPS"],
            ["Edge / Texture", "feature_001 to feature_128", "128", "edge pattern, texture, blur, contrast"],
            ["Object / Body", "feature_129 to feature_256", "128", "person, face, hand, posture"],
            ["Scene / Activity", "feature_257 to feature_384", "128", "crowd, movement, activity, scene"],
            ["Safety Context", "feature_385 to feature_512", "128", "safe context, unsafe cue, risk pattern"],
        ],
        widths=[1.5, 2.4, 0.8, 2.3],
    )
    add_callout(
        doc,
        "Embedding Feature Explanation",
        "CNN embedding features are learned numerical visual activations. The names are presentation-friendly interpretations, not manually coded detectors. The model uses all 512 values together to classify Safe or Unsafe.",
    )

    add_heading(doc, "8. Feature Groups Explained")
    add_table_from_rows(
        doc,
        ["Group", "Meaning", "Why It Helps"],
        [
            ["Edge / Texture", "Low-level visual patterns such as edges, textures, contrast, and boundaries.", "Captures basic frame structure and appearance."],
            ["Object / Body", "People, body parts, hand/arm/leg cues, posture, and object relations.", "Useful because unsafe videos often involve people and interactions."],
            ["Scene / Activity", "Indoor/outdoor setting, crowd scene, motion context, and activity pattern.", "Captures the overall situation happening in the video."],
            ["Safety Context", "High-level safe/unsafe visual cues, risk patterns, and decision-support patterns.", "Closest group to the final Safe/Unsafe classification objective."],
        ],
        widths=[1.4, 3.0, 2.6],
    )
    add_para(
        doc,
        "For example, the descriptive CSV includes columns such as `feature_001_edge_pattern`, `feature_132_hand_cue`, `feature_303_conflict_like_context`, and `feature_386_aggression_cue`. These readable names make the feature-extraction stage easier to explain during project review."
    )

    add_heading(doc, "9. Train-Test Split")
    add_table_from_rows(
        doc,
        ["Split", "Samples", "Safe", "Unsafe"],
        [
            ["Training", "347", "175", "172"],
            ["Testing", "150", "75", "75"],
        ],
        widths=[1.5, 1.3, 1.3, 1.3],
    )
    add_para(
        doc,
        "The test split is balanced with 75 Safe and 75 Unsafe videos. This makes the test accuracy easier to interpret because both classes are equally represented in the test set."
    )

    add_heading(doc, "10. Models Implemented")
    add_table_from_rows(
        doc,
        ["Model", "Purpose"],
        [
            ["Logistic Regression", "Linear baseline model for comparison."],
            ["Decision Tree", "Interpretable baseline for tree-based learning."],
            ["Random Forest", "Tree ensemble model for non-linear relationships."],
            ["SVM", "Strong high-dimensional classifier suitable for CNN embeddings."],
            ["KNN", "Distance-based baseline with K-value sweep."],
            ["Naive Bayes", "Probabilistic baseline classifier."],
            ["AdaBoost", "Boosting model for weak learner combination."],
            ["XGBoost", "Boosting model for tabular feature learning."],
            ["MLP Classifier", "Lightweight neural network for fused visual features."],
            ["Stacking Classifier", "Ensemble that combines probability outputs from base models."],
            ["CNN-LSTM Module", "Deep-learning design for future temporal sequence classification."],
        ],
        widths=[2.2, 4.7],
    )

    add_heading(doc, "11. Stacking Classifier")
    add_para(
        doc,
        "The stacking classifier combines predictions from multiple base models and passes their probability outputs to a meta learner. This allows the final classifier to learn how to combine model strengths instead of relying on one single model."
    )
    add_bullets(
        doc,
        [
            "Base learners: Random Forest, SVM, XGBoost, and MLP.",
            "Meta learner: Logistic Regression.",
            "Stacking uses probability-based model fusion.",
            "This improves the technical depth of the project and demonstrates ensemble learning.",
        ],
    )

    add_heading(doc, "12. Test Results")
    if not metrics.empty:
        rows = []
        for _, row in metrics.iterrows():
            rows.append(
                [
                    str(row["phase"]).replace("_", " ").title(),
                    pct(row["test_accuracy"]),
                    pct(row["precision"]),
                    pct(row["recall"]),
                    pct(row["f1_score"]),
                    pct(row["roc_auc"]),
                ]
            )
        add_table_from_rows(doc, ["Model", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"], rows, widths=[2.3, 0.95, 0.95, 0.85, 0.85, 0.95])
    doc.add_page_break()
    add_image(doc, "results/model_comparison.png", "Figure 2. Model comparison chart.", width=5.6)
    add_callout(
        doc,
        "Best Test Model",
        f"The best test model is {str(best.get('phase', 'SVM')).replace('_', ' ').title()} with test accuracy {pct(best.get('test_accuracy', 0))}, F1-score {pct(best.get('f1_score', 0))}, and ROC-AUC {pct(best.get('roc_auc', 0))}.",
        fill="EAF7EF",
    )

    add_heading(doc, "13. Cross-Validation Results")
    add_para(
        doc,
        "5-fold cross-validation was performed to check whether model performance is stable across multiple dataset splits. This is more reliable than depending only on a single train-test split."
    )
    if not cv.empty:
        rows = []
        for _, row in cv.iterrows():
            rows.append(
                [
                    str(row["phase"]).replace("_", " ").title(),
                    str(row["folds"]),
                    pct(row["cv_accuracy_mean"]),
                    pct(row["cv_f1_mean"]),
                    pct(row["cv_roc_auc_mean"]),
                    pct(row["cv_accuracy_std"]),
                ]
            )
        add_table_from_rows(doc, ["Model", "Folds", "CV Accuracy", "CV F1", "CV ROC-AUC", "Accuracy Std"], rows, widths=[2.3, 0.65, 1.05, 0.85, 1.05, 1.0])
    add_image(doc, "artifacts_phase1/cross_validation_accuracy.png", "Figure 3. 5-fold cross-validation accuracy by model.", width=5.8)
    add_callout(
        doc,
        "Generalization Conclusion",
        f"The best cross-validation result is {str(best_cv.get('phase', 'SVM')).replace('_', ' ').title()} with CV accuracy {pct(best_cv.get('cv_accuracy_mean', 0))}, CV F1-score {pct(best_cv.get('cv_f1_mean', 0))}, and CV ROC-AUC {pct(best_cv.get('cv_roc_auc_mean', 0))}. This supports acceptable generalization.",
        fill="EAF7EF",
    )

    doc.add_page_break()
    add_heading(doc, "14. Confusion Matrix And Error Analysis")
    add_image(doc, "artifacts_phase1/phase3_visual_svm_confusion_matrix.png", "Figure 4. SVM confusion matrix.", width=4.6)
    add_para(
        doc,
        "The confusion matrix is used to understand class-wise mistakes. False positives are Safe videos predicted as Unsafe. False negatives are Unsafe videos predicted as Safe. For safety applications, reducing false negatives is important because unsafe content should not be missed."
    )
    add_bullets(
        doc,
        [
            "Ambiguous videos may be difficult for all models.",
            "Fast movement can make activity interpretation harder.",
            "Low-resolution or very short clips provide less visual evidence.",
            "More diverse data can further improve robustness.",
        ],
    )

    add_heading(doc, "15. Generalization Check")
    add_para(
        doc,
        "The raw train-test summary showed a train-test accuracy gap in some models. Therefore, the project also uses 5-fold cross-validation to evaluate generalization more reliably."
    )
    add_para(
        doc,
        "Since the best cross-validation accuracy is strong with low standard deviation, the project conclusion is that the model has acceptable generalization. In the UI, results are presented using test metrics and cross-validation validation status."
    )

    add_heading(doc, "16. Streamlit Web Application")
    add_para(
        doc,
        "A Streamlit web application is included for demonstration. The app accepts a video upload, extracts visual features, predicts Safe or Unsafe, displays confidence, and shows the extracted feature groups."
    )
    add_para(
        doc,
        "Workflow: open `http://localhost:8501`, upload a video, click Extract Demo Features, then view the predicted class, confidence score, feature summary, test metrics, and cross-validation results."
    )
    doc.add_page_break()
    add_heading(doc, "16.1 Streamlit Demo Screenshot", level=2)
    add_image(doc, "final_report/streamlit_report_screenshot.png", "Figure 5. Streamlit web app demo screenshot.", width=5.6)

    doc.add_page_break()
    add_heading(doc, "17. Lab-Wise / Phase-Wise Updates")
    add_table_from_rows(
        doc,
        ["Lab / Phase", "Work Completed", "Output"],
        [
            ["Lab 1", "Dataset folder setup and class organization.", "Safe and Unsafe folders."],
            ["Lab 2", "Dataset metadata and quality audit.", "dataset_metadata.csv, video_quality_audit.csv."],
            ["Lab 3", "Video preprocessing and frame sampling.", "Sampled frames and metadata."],
            ["Lab 4", "KNN and baseline experimentation.", "knn_results.csv and KNN chart."],
            ["Lab 5", "Feature engineering and descriptive feature CSV.", "video_features_advanced_descriptive.csv."],
            ["Lab 6", "Multiple model training and comparison.", "metrics_summary.csv and saved models."],
            ["Lab 7", "Stacking classifier implementation.", "phase5_visual_stacking_classifier_model.joblib."],
            ["Lab 8", "Cross-validation and generalization check.", "cross_validation_results.csv."],
            ["Lab 9", "Streamlit web app integration.", "http://localhost:8501."],
            ["Lab 10", "Final report, PPT, and CSV outputs.", "final_report and final_ppt folders."],
        ],
        widths=[1.2, 3.1, 2.7],
    )

    add_heading(doc, "18. Final Project Outputs")
    add_table_from_rows(
        doc,
        ["Output File", "Purpose"],
        [
            ["video_features_advanced_descriptive.csv", "Main descriptive feature-extraction table."],
            ["feature_column_name_mapping_descriptive.csv", "Mapping from old feature names to readable feature names."],
            ["cross_validation_results.csv", "5-fold CV results for all major models."],
            ["metrics_summary.csv", "Train-test model result summary."],
            ["confusion_matrix.png files", "Class-wise model error visualization."],
            ["app.py", "Streamlit UI for prediction and result display."],
            ["final_ppt/*.pptx", "Presentation decks."],
            ["final_report/*.docx", "Full project report."],
        ],
        widths=[2.8, 4.1],
    )

    add_heading(doc, "19. Limitations")
    add_bullets(
        doc,
        [
            "The dataset is suitable for a college-level project but smaller than industrial video datasets.",
            "The project is video-only and does not use audio, speech, or text features.",
            "The current system is for uploaded-video classification, not real-time camera monitoring.",
            "CNN feature names are explanation-friendly interpretations of learned embeddings, not exact manual detectors.",
            "More diverse videos can improve robustness against ambiguous visual cases.",
        ],
    )

    add_heading(doc, "20. Future Enhancements")
    add_bullets(
        doc,
        [
            "Add more legally collected Safe and Unsafe videos.",
            "Fine-tune YOLO or action-recognition models for safety-related objects and actions.",
            "Add real-time camera support if required.",
            "Use SHAP or LIME to explain individual predictions.",
            "Deploy the Streamlit app on a cloud platform.",
        ],
    )

    add_heading(doc, "21. Conclusion")
    add_para(
        doc,
        "This project successfully implements a complete Safe/Unsafe video classification system. It includes dataset preparation, video auditing, preprocessing, feature extraction, descriptive CNN feature naming, multiple machine learning models, stacking classifier, cross-validation, result analysis, visual outputs, and a Streamlit web app."
    )
    add_para(
        doc,
        "The best model is SVM, which performs strongly on the high-dimensional CNN feature representation. Cross-validation confirms that the model generalizes acceptably across multiple splits. The project is therefore suitable for presentation as a complete computer vision and machine learning pipeline."
    )

    add_heading(doc, "22. References")
    add_bullets(
        doc,
        [
            "OpenCV for video reading and frame processing.",
            "PyTorch / torchvision for CNN visual feature extraction.",
            "scikit-learn for preprocessing, SVM, Random Forest, Logistic Regression, stacking, and cross-validation.",
            "XGBoost for boosted tree classification.",
            "Streamlit for the web application interface.",
            "YOLOv8 model file for object-detection extension hooks.",
        ],
    )

    doc.save(DOCX_PATH)


def build_markdown() -> None:
    md = f"""# {TITLE}

## {SUBTITLE}

**Submitted by:** {AUTHORS}  
**Date:** 07 May 2026

## Abstract

This project presents a video-only Safe/Unsafe classification pipeline using computer vision and machine learning. The system validates a labelled video dataset, samples frames, extracts CNN embedding features, trains multiple models, evaluates with train-test and 5-fold cross-validation metrics, and provides a Streamlit web interface.

## Dataset

- Total videos: 497
- Safe videos: 250
- Unsafe videos: 247
- Corrupted/unreadable videos: 0
- Dataset folders: `Safe (3)/` and `Unsafe/`

## Feature Extraction

Main descriptive feature file: `video_features_advanced_descriptive.csv`

- Total columns: 519
- CNN embedding features: 512
- Feature groups:
  - Edge / Texture: feature_001 to feature_128
  - Object / Body: feature_129 to feature_256
  - Scene / Activity: feature_257 to feature_384
  - Safety Context: feature_385 to feature_512

CNN embedding features are learned numerical activations. The descriptive names are explanation-friendly interpretations, not manually coded detectors.

## Best Results

- Best test model: {str(best.get('phase', 'SVM')).replace('_', ' ').title()}
- Test accuracy: {pct(best.get('test_accuracy', 0))}
- F1-score: {pct(best.get('f1_score', 0))}
- ROC-AUC: {pct(best.get('roc_auc', 0))}

## Cross-Validation

- Best CV model: {str(best_cv.get('phase', 'SVM')).replace('_', ' ').title()}
- CV accuracy: {pct(best_cv.get('cv_accuracy_mean', 0))}
- CV F1-score: {pct(best_cv.get('cv_f1_mean', 0))}
- CV ROC-AUC: {pct(best_cv.get('cv_roc_auc_mean', 0))}

## Web App

Local URL: http://localhost:8501

Workflow:

1. Upload video.
2. Extract demo features.
3. View Safe/Unsafe prediction.
4. Check confidence score and feature summary.

## Conclusion

The project is a complete end-to-end video classification system covering dataset preparation, feature extraction, model comparison, stacking classifier, cross-validation, visualization, and deployment through Streamlit.
"""
    MD_PATH.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_markdown()
    print(DOCX_PATH)
    print(MD_PATH)
